#!/usr/bin/env python3
"""Word .docx → HTML 语义标签转换器（python-docx）
输出全语义化 HTML：标题→<h1~h4>、正文→<p>、着重号→<span class="emphasis-dot">、下划线→<u>
Tiptap 解析时保留元素标签（丢弃 class），主题 CSS 通过元素选择器精准命中。
"""
import sys, json, os, re, html as html_mod, base64
from collections import Counter
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Emu


# ── 对齐映射 ──
ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.CENTER: 'center',
    WD_ALIGN_PARAGRAPH.RIGHT: 'right',
    WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
}

# ── 样式名 → 标题级别映射（中/英文） ──
HEADING_MAP = {
    'heading 1': 2, '标题 1': 2,
    'heading 2': 3, '标题 2': 3,
    'heading 3': 4, '标题 3': 4,
    'heading 4': 4, '标题 4': 4,
    'title': 1, '标题': 1,
}


def get_heading_level(style_name):
    """通过段落样式名推断标题级别"""
    low = style_name.lower() if style_name else ''
    for key, level in HEADING_MAP.items():
        if key in low:
            return level
    return None


def analyze_body_size(paragraphs):
    """文档级分析：扫描所有段落，用字号中位数作为正文基准字号。
    
    不同老师/学科的文档正文字号不同（12pt/14pt/小四等），
    死阈值（如 ≥16pt=标题）无法适应所有文档。通过中位数自动检测
    该文档的正文标准字号，标题判定使用相对差异（≥body+4pt=H1）。
    """
    sizes = []
    for para in paragraphs:
        text = para.text.strip()
        if not text or len(text) < 2:
            continue
        if para.runs and para.runs[0].font.size:
            sizes.append(para.runs[0].font.size.pt)
    if not sizes:
        return 12  # 默认12pt
    # 中位数：比均值更能抵抗极端字号（如 22pt 标题）的干扰
    sorted_sizes = sorted(sizes)
    result = sorted_sizes[len(sorted_sizes) // 2]
    print(f'[word_to_html] 正文字号基准(body_size)={result}pt （{len(sizes)}个段落，范围{sorted_sizes[0]}-{sorted_sizes[-1]}pt）', file=sys.stderr)
    return result


def looks_like_heading(para, body_size=12):
    """智能标题分类：相对字号 + 视觉特征 + 内容分析（非纯正则）
    
    设计理念：
      - 相对字号：不以死阈值（≥16pt）判定，而是看比正文大多少
      - 对齐感知：主标题必须居中，章节标题可左可居中
      - 内容判断：句末标点（。！？）→正文特征，非标题
      - 有意义的短文本才算标题
    
    参数：
      body_size — 文档正文字号（中位数），默认12pt
    返回 (level, is_heading) — level 1-4"""
    if not para.runs:
        return None, False

    first_run = para.runs[0]
    if not first_run.bold:
        return None, False

    text = para.text.strip()
    text_len = len(text)

    if text_len > 80:
        return None, False

    # 获取字号
    font_size = None
    if first_run.font.size:
        font_size = first_run.font.size.pt

    if font_size is None:
        if text_len <= 25:
            return 3, True
        return None, False

    is_centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
    size_diff = font_size - body_size

    # ── 内容判断：正文特征 ──
    # 句末标点是正文强特征（标题通常不以。！？结尾）
    ends_like_body = text.endswith(('。', '！', '？', '…', '.', '!', '?'))
    # 有主谓结构的长句 → 正文（简单启发：包含"的"/"是"/"在"等虚词+较长）
    has_body_structure = text_len > 35 and any(w in text for w in ['的', '是', '在', '了', '着'])

    # ── 日志：视觉检测决策过程 ──
    reason_parts = [f'bold ✓', f'size={font_size}pt(body={body_size},diff={size_diff:+d}pt)']
    if is_centered:
        reason_parts.append('居中 ✓')
    if ends_like_body:
        reason_parts.append('句末标点 ✗')
    if has_body_structure:
        reason_parts.append('正文句式 ✗')

    # ── 标题分级（相对字号 + 对齐 + 内容特征） ──

    # H1：比正文大4pt以上 + 居中 + 粗体 + 短文本（≤50字 + 非正文句）
    if size_diff >= 4 and is_centered and text_len <= 50 and not ends_like_body:
        print(f'[word_to_html] H1 ← "{text[:30]}" ({chr(44).join(reason_parts)})', file=sys.stderr)
        return 1, True

    # H1 特例：极大字（≥6pt差），即使不居中（文档第一个大字标题）
    if size_diff >= 6 and text_len <= 50 and not ends_like_body:
        print(f'[word_to_html] H1(特例) ← "{text[:30]}" ({chr(44).join(reason_parts)})', file=sys.stderr)
        return 1, True

    # H2：比正文大2pt以上 + 粗体 + ≤80字，非纯正文句
    if size_diff >= 2 and not (ends_like_body and has_body_structure):
        print(f'[word_to_html] H2 ← "{text[:30]}" ({chr(44).join(reason_parts)})', file=sys.stderr)
        return 2, True

    # H3：≥正文或略大于正文 + 粗体 + (短文本 或 居中)
    if is_bold_heading_like(text_len, is_centered, ends_like_body):
        print(f'[word_to_html] H3 ← "{text[:30]}" ({chr(44).join(reason_parts)})', file=sys.stderr)
        return 3, True

    # H4：正文字号粗体 + 极短（≤20字）+ 非正文句
    if size_diff >= 0 and text_len <= 20 and not ends_like_body:
        print(f'[word_to_html] H4 ← "{text[:30]}" ({chr(44).join(reason_parts)})', file=sys.stderr)
        return 4, True

    # 粗体但不符合任何标题规则 → 标记为正文（让 LLM 最后判断）
    if ends_like_body or text_len > 80:
        print(f'[word_to_html] P(粗体但不像标题) ← "{text[:30]}" ({chr(44).join(reason_parts)})', file=sys.stderr)
    return None, False


def is_bold_heading_like(text_len, is_centered, ends_like_body):
    """判断是否像子标题：短文本 或 居中，且不像正文句"""
    if ends_like_body:
        return False
    if text_len <= 40:
        return True
    if is_centered and text_len <= 50:
        return True
    return False


def get_emphasis_dot(run):
    """检测 Word 着重号（加点字）: w:rPr/w:em[@w:val='dot']"""
    try:
        rpr = run._element.find(qn('w:rPr'))
        if rpr is not None:
            em = rpr.find(qn('w:em'))
            if em is not None and em.get(qn('w:val')) == 'dot':
                return True
    except Exception:
        pass
    return False


def emu_to_pt(emu_val):
    """EMU → pt"""
    if emu_val is None:
        return None
    return round(emu_val / 12700, 1)


def extract_images(doc):
    """提取文档中所有图片，返回 {rId: base64_data_uri}"""
    images = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                ext = os.path.splitext(rel.target_part.partname)[-1].lower()
                mime_map = {'.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                            '.gif': 'image/gif', '.bmp': 'image/bmp', '.webp': 'image/webp'}
                mime = mime_map.get(ext, 'image/png')
                b64 = base64.b64encode(img_data).decode('ascii')
                images[rel.rId] = f'data:{mime};base64,{b64}'
            except Exception:
                pass
    return images


def get_run_html(run, doc_images):
    """将 Word run 转换为行内 HTML 标签（语义化 + style 兜底）"""
    text = html_mod.escape(run.text or '')
    if not text and not run._element.findall(qn('w:drawing')):
        return ''

    tags_open = []
    tags_close = []
    styles = []

    # 粗体
    if run.bold:
        tags_open.append('<strong>')
        tags_close.insert(0, '</strong>')

    # 斜体
    if run.italic:
        tags_open.append('<em>')
        tags_close.insert(0, '</em>')

    # 下划线
    if run.underline:
        tags_open.append('<u>')
        tags_close.insert(0, '</u>')

    # 删除线
    if run.font.strike:
        tags_open.append('<s>')
        tags_close.insert(0, '</s>')

    # ⭐ 着重号（加点字）→ Tiptap EmphasisDot mark 解析 <span class="emphasis-dot">
    if get_emphasis_dot(run):
        tags_open.append('<span class="emphasis-dot">')
        tags_close.insert(0, '</span>')

    # 上标/下标
    if run.font.superscript:
        tags_open.append('<sup>')
        tags_close.insert(0, '</sup>')
    elif run.font.subscript:
        tags_open.append('<sub>')
        tags_close.insert(0, '</sub>')

    # 字号
    if run.font.size:
        styles.append(f'font-size:{run.font.size.pt}pt')

    # 颜色
    if run.font.color and run.font.color.rgb:
        styles.append(f'color:#{run.font.color.rgb}')

    # 字体名
    if run.font.name:
        styles.append(f"font-family:'{run.font.name}'")

    # 高亮（底纹）
    try:
        rpr = run._element.find(qn('w:rPr'))
        if rpr is not None:
            highlight = rpr.find(qn('w:highlight'))
            if highlight is not None:
                color = highlight.get(qn('w:val'))
                if color and color != 'none':
                    styles.append(f'background-color:#{color}')
    except Exception:
        pass

    # 行内图片
    drawings = run._element.findall(qn('w:drawing'))
    image_html = ''
    for drawing in drawings:
        blip = drawing.find('.//' + qn('a:blip'))
        if blip is not None:
            embed = blip.get(qn('r:embed'))
            if embed and embed in doc_images:
                image_html += f'<img src="{doc_images[embed]}" style="max-width:100%;height:auto">'

    inner = text or '&#8203;'  # 零宽空格兜底空文本
    if styles:
        inner = f'<span style="{";".join(styles)}">{inner}</span>'

    return ''.join(tags_open) + image_html + inner + ''.join(tags_close)


def is_list_item(para):
    """检测段落是否为列表项"""
    try:
        ppr = para._element.find(qn('w:pPr'))
        if ppr is not None:
            numPr = ppr.find(qn('w:numPr'))
            return numPr is not None
    except Exception:
        pass
    return False


def get_paragraph_html(para, doc_images, body_size=12):
    """将 Word 段落转换为语义化 HTML"""
    style_name = para.style.name if para.style else ''

    # 检测标题级别（① 样式名 → ② 智能视觉检测）
    heading_level = get_heading_level(style_name)
    if not heading_level:
        visual_level, is_heading = looks_like_heading(para, body_size)
        if is_heading:
            heading_level = visual_level

    # 段落级样式
    para_styles = []

    # 对齐
    if para.alignment in ALIGN_MAP:
        para_styles.append(f'text-align:{ALIGN_MAP[para.alignment]}')

    # 首行缩进
    indent_pt = emu_to_pt(para.paragraph_format.first_line_indent)
    if indent_pt and indent_pt > 0:
        para_styles.append(f'text-indent:{indent_pt}pt')
    elif not heading_level:
        # 正文默认两个字符缩进（≈ 2em，编辑器的 INDENT_STEP=24px）
        para_styles.append('text-indent:2em')

    # 构建行内内容
    parts = [get_run_html(run, doc_images) for run in para.runs]
    content = ''.join(p for p in parts if p)

    # 空段落
    if not content.strip():
        return '<p><br></p>' if not heading_level else ''

    # 输出语义标签（h1~h4 或 p）
    # 🔑 标题段落级加粗：冗余写入确保 JS extractVisualFeatures 能检测到
    #    （即使 <strong> 子标签被过滤，paragraph style 中也有 bold 信号）
    if heading_level:
        para_styles.insert(0, 'font-weight:bold')
    style_str = f' style="{";".join(para_styles)}"' if para_styles else ''
    if heading_level:
        tag = f'h{heading_level}'
        return f'<{tag}{style_str}>{content}</{tag}>'
    else:
        return f'<p{style_str}>{content}</p>'


def convert_docx_to_html(file_path):
    """主转换函数：.docx → HTML body 内容（按文档顺序保留表格位置）"""
    if not os.path.exists(file_path):
        return {'error': f'文件不存在: {file_path}'}

    try:
        doc = Document(file_path)
        doc_images = extract_images(doc)
        html_parts = []

        # 🔧 第一遍：文档级字号分析 — 自动检测正文字号（中位数）
        body_size = analyze_body_size(doc.paragraphs)

        # 🔧 按文档原始顺序遍历：段落(w:p)和表格(w:tbl)交错处理
        para_map = {p._element: p for p in doc.paragraphs}
        table_map = {t._tbl: t for t in doc.tables}

        in_list = False

        for child in doc.element.body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':  # 段落
                para = para_map.get(child)
                if para is None:
                    continue

                is_item = is_list_item(para)

                if is_item and not in_list:
                    html_parts.append('<ul>')
                    in_list = True

                if is_item:
                    content = ''.join(get_run_html(run, doc_images) for run in para.runs)
                    html_parts.append(f'<li>{content}</li>')
                else:
                    if in_list:
                        html_parts.append('</ul>')
                        in_list = False
                    html_parts.append(get_paragraph_html(para, doc_images, body_size))

            elif tag == 'tbl':  # 表格
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False

                table = table_map.get(child)
                if table is None:
                    continue

                rows = []
                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        cell_html = ''.join(
                            get_run_html(run, doc_images)
                            for p in cell.paragraphs
                            for run in p.runs
                        )
                        cells.append(f'<td>{cell_html}</td>')
                    rows.append(f'<tr>{"".join(cells)}</tr>')
                html_parts.append(f'<table>{"".join(rows)}</table>')

        if in_list:
            html_parts.append('</ul>')

        return {'html': '\n'.join(html_parts)}

    except Exception as e:
        return {'error': str(e)}


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(json.dumps({'error': '缺少文件路径参数'}, ensure_ascii=False))
        sys.exit(1)

    result = convert_docx_to_html(sys.argv[1])
    print(json.dumps(result, ensure_ascii=False))
